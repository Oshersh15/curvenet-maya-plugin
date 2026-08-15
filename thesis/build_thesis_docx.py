import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "thesis_content.md"
OUTPUT = ROOT / "Curvenet_MSc_Thesis_Draft.docx"

NAVY = "17233B"
TEAL = "168A8A"
MINT = "DDEEEE"
LIGHT = "F2F4F7"
MID = "6B7280"
DARK = "20242A"
ORANGE = "D97706"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "Right-click and choose Update Field to generate the contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, fallback, fld_end])


def enable_update_fields(document):
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_run_font(run, name="Aptos", size=10.5, color=DARK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_rich_text(paragraph, text, base_size=10.5, base_color=DARK):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, color=base_color, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Menlo", size=base_size - 0.5, color=NAVY)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=base_color)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 16, TEAL, 0, 12),
        ("Heading 1", 21, NAVY, 6, 10),
        ("Heading 2", 14, TEAL, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = "Cambria Math"
    equation.font.size = Pt(11)
    equation.font.color.rgb = RGBColor.from_string(NAVY)
    equation.paragraph_format.space_before = Pt(7)
    equation.paragraph_format.space_after = Pt(7)
    equation.paragraph_format.keep_together = True


def configure_section(section, first=False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    if not first:
        header = section.header
        p = header.paragraphs[0]
        p.text = "CURVENET  /  MSc THESIS DRAFT"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            set_run_font(run, size=7.8, color=MID, bold=True)
        p.paragraph_format.space_after = Pt(0)
        add_page_number(section.footer.paragraphs[0])


def add_chapter_heading(document, text, level):
    if level == 1 and (
        re.match(r"\d+\.", text)
        or text in {"References"}
        or text.startswith("Appendix")
    ):
        document.add_page_break()
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), TEAL)
        border.append(bottom)
        pPr.append(border)
    return p


def add_placeholder(document, instruction, caption):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    set_cell_margins(cell, 170, 190, 170, 190)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AUTHOR FIGURE ACTION\n")
    set_run_font(r, size=9, color=ORANGE, bold=True)
    r = p.add_run(instruction)
    set_run_font(r, size=9.5, color=DARK)
    p.paragraph_format.space_after = Pt(0)
    cap = document.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Planned figure - " + caption)


def add_image(document, path, caption, figure_number):
    image_path = Path(path)
    if not image_path.exists():
        add_placeholder(document, f"Image file not found: {path}", caption)
        return figure_number + 1
    with Image.open(image_path) as im:
        width_px, height_px = im.size
    max_width = 6.25
    max_height = 7.2
    ratio = width_px / max(height_px, 1)
    width = min(max_width, max_height * ratio)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = document.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Figure {figure_number}. {caption}")
    return figure_number + 1


def add_equation(document, expression, equation_number):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(14.2)
    table.columns[1].width = Cm(1.2)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, 70, 80, 70, 80)
    p = left.paragraphs[0]
    p.style = document.styles["Equation"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(p, expression, base_size=11, base_color=NAVY)
    q = right.paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = q.add_run(f"({equation_number})")
    set_run_font(r, size=9, color=MID)
    return equation_number + 1


def add_markdown_table(document, lines):
    parsed = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        parsed.append(cells)
    headers = parsed[0]
    rows = [row for row in parsed[2:] if row]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, value in enumerate(headers):
        set_cell_shading(header_cells[index], NAVY)
        p = header_cells[index].paragraphs[0]
        add_rich_text(p, value, base_size=9, base_color="FFFFFF")
        for run in p.runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[index], LIGHT)
            p = cells[index].paragraphs[0]
            add_rich_text(p, value, base_size=9)
            if index > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(document, lines):
    title = lines[0][2:].strip()
    subtitle = lines[2][3:].strip()
    author_lines = [line for line in lines[4:9] if line.strip()]
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("MSc PROJECT  /  2026")
    set_run_font(r, size=9, color=TEAL, bold=True)
    p = document.add_paragraph(style="Title")
    p.add_run(title)
    p = document.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    accent = document.add_table(rows=1, cols=1)
    accent.autofit = False
    accent.columns[0].width = Cm(4.2)
    cell = accent.cell(0, 0)
    set_cell_shading(cell, TEAL)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].add_run(" ")
    for line in author_lines:
        if not line.strip():
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        text = line.replace("**", "").replace("  ", "").strip()
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK, bold=text.startswith("Osher"))


def build():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0], first=True)
    enable_update_fields(document)

    # Cover content through supervisor line, then the cover image marker.
    add_cover(document, lines[:10])
    figure_number = 1
    equation_number = 1
    first_image = re.match(r"\[\[IMAGE:(.*?)\|(.*?)\]\]", lines[10])
    if first_image:
        figure_number = add_image(
            document, first_image.group(1), first_image.group(2), figure_number
        )
    document.add_page_break()

    # Front-matter section starts after the cover image.
    index = 11
    buffer = []
    in_quote = False
    contents_inserted = False

    def flush_buffer():
        nonlocal buffer
        if not buffer:
            return
        paragraph_text = " ".join(part.strip() for part in buffer).strip()
        buffer = []
        if paragraph_text:
            p = document.add_paragraph()
            add_rich_text(p, paragraph_text)

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_buffer()
            index += 1
            continue

        image_match = re.fullmatch(r"\[\[IMAGE:(.*?)\|(.*?)\]\]", line)
        figure_match = re.fullmatch(r"\[\[FIGURE:(.*?)\|(.*?)\]\]", line)
        equation_match = re.fullmatch(r"\[\[EQUATION:(.*?)\]\]", line)
        if image_match:
            flush_buffer()
            figure_number = add_image(
                document, image_match.group(1), image_match.group(2), figure_number
            )
        elif figure_match:
            flush_buffer()
            add_placeholder(document, figure_match.group(1), figure_match.group(2))
            figure_number += 1
        elif equation_match:
            flush_buffer()
            equation_number = add_equation(
                document, equation_match.group(1), equation_number
            )
        elif line.startswith("### "):
            flush_buffer()
            add_chapter_heading(document, line[4:].strip(), 2)
        elif line.startswith("## "):
            flush_buffer()
            heading = line[3:].strip()
            if heading == "1. Introduction" and not contents_inserted:
                document.add_page_break()
                contents_heading = document.add_paragraph(style="Heading 1")
                contents_heading.add_run("Contents")
                add_toc(document.add_paragraph())
                contents_inserted = True
            if heading == "Abstract":
                add_chapter_heading(document, heading, 1)
            elif heading == "Acknowledgements":
                add_chapter_heading(document, heading, 1)
            elif heading == "Declaration of Generative AI Use":
                add_chapter_heading(document, heading, 1)
            else:
                add_chapter_heading(document, heading, 1)
        elif line.startswith("# "):
            flush_buffer()
            add_chapter_heading(document, line[2:].strip(), 1)
        elif line.startswith("> "):
            flush_buffer()
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            add_rich_text(p, line[2:].strip(), base_size=11, base_color=NAVY)
            for run in p.runs:
                run.italic = True
        elif re.match(r"^\d+\. ", line):
            flush_buffer()
            p = document.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            flush_buffer()
            p = document.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
        elif (
            line.startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip())
        ):
            flush_buffer()
            table_lines = [line, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_lines)
            continue
        elif line.startswith("[AUTHOR ACTION"):
            flush_buffer()
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, "FFF7E8")
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(line)
            set_run_font(r, size=9.5, color=ORANGE, bold=True)
        elif line.startswith("**") and line.endswith("**") and len(line) < 120:
            flush_buffer()
            p = document.add_paragraph()
            add_rich_text(p, line, base_size=10.5, base_color=NAVY)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
        else:
            buffer.append(line)
        index += 1
    flush_buffer()

    # Apply body headers and page numbers to the existing section.
    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.text = "CURVENET  /  MSc THESIS DRAFT"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        set_run_font(run, size=7.8, color=MID, bold=True)
    add_page_number(section.footer.paragraphs[0])

    document.core_properties.title = (
        "Topology-Independent Character Deformation Using Curvenets"
    )
    document.core_properties.subject = "MSc Thesis Draft"
    document.core_properties.author = "Osher [Surname]"
    document.core_properties.keywords = (
        "Curvenet, Maya, topology-independent deformation, profile curves"
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
